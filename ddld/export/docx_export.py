"""
Export the full run to Word (.docx): disclaimer, verdict tally, a color-coded
claims table, and the raw transcript. Requires python-docx.
"""
from __future__ import annotations

from collections import Counter

from ..types import CheckedClaim, Utterance, Verdict


def export_docx(
    path: str,
    transcript: list[Utterance],
    checked: list[CheckedClaim],
    disclaimer: str,
    title: str = "Debate Lie Detector — Session Log",
) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(title, level=0)

    # Disclaimer up top (spec requirement).
    d = doc.add_paragraph()
    run = d.add_run("DISCLAIMER: " + disclaimer)
    run.italic = True
    run.font.size = Pt(9)

    # Tally.
    tally = Counter(c.verdict for c in checked)
    doc.add_heading("Summary", level=1)
    summary = doc.add_paragraph()
    for v in (Verdict.SUPPORTED, Verdict.CONTRADICTED, Verdict.DISPUTED, Verdict.PENDING):
        summary.add_run(f"{v.label}: {tally.get(v, 0)}    ")

    # Claims table.
    doc.add_heading("Checked claims", level=1)
    cols = ["Time", "Speaker", "Claim", "Verdict", "Conf.", "Reasoning", "Sources"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, name in enumerate(cols):
        table.rows[0].cells[i].paragraphs[0].add_run(name).bold = True

    for c in checked:
        row = c.as_row()
        cells = table.add_row().cells
        cells[0].text = f"{row['timestamp_s']}s"
        cells[1].text = row["speaker"]
        cells[2].text = row["claim"]
        # Verdict cell: colored bold label.
        vp = cells[3].paragraphs[0]
        vrun = vp.add_run(c.verdict.label)
        vrun.bold = True
        vrun.font.color.rgb = _rgb(c.verdict.hex)
        cells[4].text = f"{row['confidence']:.2f}" if c.verdict is not Verdict.PENDING else "—"
        cells[5].text = row["reasoning"]
        srccell = cells[6]
        for s in c.sources:
            p = srccell.add_paragraph()
            p.add_run(f"[{s.stance}] {s.url}").font.size = Pt(8)

    # Transcript.
    doc.add_heading("Full transcript", level=1)
    for u in transcript:
        p = doc.add_paragraph()
        head = f"[{u.start:.0f}s] " + (f"{u.speaker}: " if u.speaker else "")
        p.add_run(head).bold = True
        p.add_run(u.text)

    doc.save(path)
    return path


def _rgb(hexstr: str):
    from docx.shared import RGBColor
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))
